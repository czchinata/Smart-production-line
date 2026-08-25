from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Any
import random

import numpy as np
import pandas as pd

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    Document = None


SEED = 42
random.seed(SEED)
np.random.seed(SEED)


PRODUCTS = {
    "single_crystal_silicon": {
        "name": "单晶硅",
        "route": [
            ("SC01", "凸面铣磨", ["aspherical_grinding"]),
            ("SC02", "凹面粗铣磨", ["aspherical_grinding"]),
            ("SC03", "清洗", ["cleaning"]),
            ("SC04", "凸面研磨", ["grinding_polishing"]),
            ("SC05", "清洗", ["cleaning"]),
            ("SC06", "凸面抛光", ["grinding_polishing"]),
            ("SC07", "清洗", ["cleaning"]),
            ("SC08", "定中心磨边", ["aspherical_grinding"]),
            ("SC09", "清洗", ["cleaning"]),
            ("SC10", "凹面粗车", ["diamond_turning"]),
            ("SC11", "凹面精车", ["diamond_turning"]),
            ("SC12", "清洗", ["cleaning"]),
            ("SC13", "凹面抛光", ["magnetorheological_polisher"]),
            ("SC14", "清洗", ["cleaning"]),
            ("SC15", "镀膜", ["vacuum_coater"]),
        ],
    },
    "chalcogenide_glass": {
        "name": "硫系玻璃",
        "route": [
            ("CG01", "凸面成型粗车削", ["diamond_turning"]),
            ("CG02", "凸面成型精车削", ["diamond_turning"]),
            ("CG03", "凹面粗铣磨", ["aspherical_grinding"]),
            ("CG04", "凹面精铣磨", ["aspherical_grinding"]),
            ("CG05", "凹面粗车削", ["diamond_turning"]),
            ("CG06", "凹面精车削", ["diamond_turning"]),
            ("CG07", "凸面粗车削", ["diamond_turning"]),
            ("CG08", "凸面精车削", ["diamond_turning"]),
            ("CG09", "清洗", ["cleaning"]),
            ("CG10", "镀膜", ["vacuum_coater"]),
        ],
    },
    "microcrystalline_glass": {
        "name": "微晶玻璃",
        "route": [
            ("MG01", "平面研磨", ["grinding_polishing"]),
            ("MG02", "上盘", ["grinding_polishing"]),
            ("MG03", "铣磨球面", ["aspherical_grinding"]),
            ("MG04", "粗铣非球面", ["aspherical_grinding"]),
            ("MG05", "精铣非球面", ["aspherical_grinding"]),
            ("MG06", "粗抛非球面", ["aspherical_polishing"]),
            ("MG07", "精抛非球面", ["aspherical_polishing"]),
            ("MG08", "磨外圆和端面", ["aspherical_grinding"]),
            ("MG09", "磨内孔", ["aspherical_grinding"]),
            ("MG10", "下盘", ["grinding_polishing"]),
            ("MG11", "非球面超精抛光", ["ion_beam_polisher"]),
            ("MG12", "镀膜机准备", ["vacuum_coater"]),
            ("MG13", "薄膜镀制", ["vacuum_coater"]),
        ],
    },
}


EQUIPMENT_TYPES = {
    "aspherical_grinding": {
        "name": "非球面数控铣磨机",
        "count": 2,
    },
    "aspherical_polishing": {
        "name": "非球面数控抛光机",
        "count": 2,
    },
    "ion_beam_polisher": {
        "name": "离子束抛光机",
        "count": 1,
    },
    "magnetorheological_polisher": {
        "name": "磁流变抛光机",
        "count": 1,
    },
    "diamond_turning": {
        "name": "单点金刚石车床",
        "count": 2,
    },
    "grinding_polishing": {
        "name": "研磨抛光机",
        "count": 2,
    },
    "vacuum_coater": {
        "name": "真空镀膜机",
        "count": 1,
    },
    "cleaning": {
        "name": "清洗机",
        "count": 2,
    },
}


EVENT_TYPES = [
    ("material_shortage", "生产物料缺少"),
    ("material_delay", "原材料到货延迟"),
    ("consumable_rework", "砂轮等易耗品修磨"),
    ("consumable_shortage", "辅料缺料"),
    ("plan_change", "工序计划临时调整"),
    ("quality_hold", "质量卡"),
    ("task_change", "任务计划变更"),
    ("insert_order", "生产插单"),
    ("operator_absence", "操作人员临时缺岗"),
    ("quality_stop", "产品质量问题停产"),
    ("equipment_fault", "设备故障停产"),
    ("facility_fault", "空调或基础设施故障"),
    ("maintenance", "设备维修保养"),
    ("safety_incident", "安全事故异常处理"),
]


@dataclass
class Config:
    start_time: str = "2025-01-01 00:00:00"
    hours: int = 72
    telemetry_minutes: int = 1
    workpiece_count: int = 36
    output_dir: str = "synthetic_production_input"
    random_seed: int = SEED


def equipment_master(cfg: Config) -> pd.DataFrame:
    records = []

    for equipment_type, info in EQUIPMENT_TYPES.items():
        for i in range(1, info["count"] + 1):
            records.append(
                {
                    "equipment_id": f"{equipment_type}_{i:02d}",
                    "equipment_type": equipment_type,
                    "equipment_name": info["name"],
                    "manufacturer": "synthetic",
                    "install_date": "2022-01-01",
                    "rated_power_kw": round(
                        np.random.uniform(8, 45), 2
                    ),
                    "status": "active",
                }
            )

    return pd.DataFrame(records)


def generate_work_orders(cfg: Config, master: pd.DataFrame):
    rng = np.random.default_rng(cfg.random_seed)

    start = pd.Timestamp(cfg.start_time, tz="UTC")
    records = []
    quality_records = []
    plan_records = []
    operator_records = []

    operators = [f"OP_{i:03d}" for i in range(1, 13)]

    for i in range(cfg.workpiece_count):
        product_key = list(PRODUCTS)[i % len(PRODUCTS)]
        product_name = PRODUCTS[product_key]["name"]
        route = PRODUCTS[product_key]["route"]

        work_order_id = f"WO_{i + 1:04d}"
        workpiece_id = f"WP_{i + 1:04d}"

        current_time = start + pd.Timedelta(
            minutes=int(i * 75)
        )

        operator_id = operators[i % len(operators)]

        for seq, (process_code, process_name, compatible_types) in enumerate(route, 1):
            process_type = compatible_types[0]

            candidates = master[
                master["equipment_type"] == process_type
            ]["equipment_id"].tolist()

            if not candidates:
                equipment_id = None
            else:
                equipment_id = candidates[
                    (i + seq) % len(candidates)
                ]

            duration = int(
                rng.integers(
                    10 if process_type == "cleaning" else 20,
                    35 if process_type == "cleaning" else 100,
                )
            )

            planned_start = current_time
            planned_end = current_time + pd.Timedelta(
                minutes=duration
            )

            actual_start = planned_start + pd.Timedelta(
                minutes=int(rng.integers(0, 4))
            )
            actual_end = actual_start + pd.Timedelta(
                minutes=duration + int(rng.integers(-3, 8))
            )

            process_status = "completed"

            records.append(
                {
                    "work_order_id": work_order_id,
                    "workpiece_id": workpiece_id,
                    "product_type": product_key,
                    "product_name": product_name,
                    "process_sequence": seq,
                    "process_code": process_code,
                    "process_name": process_name,
                    "equipment_type": process_type,
                    "equipment_id": equipment_id,
                    "planned_start": planned_start,
                    "planned_end": planned_end,
                    "actual_start": actual_start,
                    "actual_end": actual_end,
                    "operator_id": operator_id,
                    "process_status": process_status,
                }
            )

            plan_records.append(
                {
                    "work_order_id": work_order_id,
                    "workpiece_id": workpiece_id,
                    "product_type": product_key,
                    "process_code": process_code,
                    "process_sequence": seq,
                    "planned_start": planned_start,
                    "planned_end": planned_end,
                    "plan_version": 1,
                    "plan_change_flag": 0,
                    "quality_hold_flag": 0,
                    "expedite_flag": 0,
                    "insert_order_flag": 0,
                }
            )

            if process_type != "cleaning":
                quality_records.append(
                    {
                        "work_order_id": work_order_id,
                        "workpiece_id": workpiece_id,
                        "product_type": product_key,
                        "process_code": process_code,
                        "inspection_time": actual_end,
                        "dimension_error": round(
                            rng.normal(0, 0.8), 4
                        ),
                        "surface_roughness_ra": round(
                            abs(rng.normal(0.12, 0.04)), 4
                        ),
                        "curvature_error": round(
                            rng.normal(0, 0.5), 4
                        ),
                        "thickness_error": round(
                            rng.normal(0, 0.3), 4
                        ),
                        "quality_result": (
                            "pass"
                            if rng.random() > 0.08
                            else "rework"
                        ),
                        "quality_card_flag": 0,
                        "rework_flag": 0,
                        "scrap_flag": 0,
                    }
                )

            current_time = actual_end + pd.Timedelta(
                minutes=int(rng.integers(2, 8))
            )

        operator_records.append(
            {
                "operator_id": operator_id,
                "work_order_id": work_order_id,
                "workpiece_id": workpiece_id,
                "shift_id": f"SHIFT_{(i % 3) + 1}",
                "skill_level": int(rng.integers(1, 4)),
                "available_flag": 1,
                "absence_flag": 0,
                "assigned_process": "multi_process",
            }
        )

    return (
        pd.DataFrame(records),
        pd.DataFrame(quality_records),
        pd.DataFrame(plan_records),
        pd.DataFrame(operator_records),
    )


def generate_events(
    cfg: Config,
    equipment: pd.DataFrame,
    production: pd.DataFrame,
) -> pd.DataFrame:
    rng = np.random.default_rng(cfg.random_seed + 100)

    start = pd.Timestamp(cfg.start_time, tz="UTC")
    end = start + pd.Timedelta(hours=cfg.hours)

    rows = []

    for event_id, (event_type, event_name) in enumerate(
        EVENT_TYPES, 1
    ):
        event_start = start + pd.Timedelta(
            minutes=int(rng.integers(120, cfg.hours * 60 - 180))
        )
        duration = int(rng.integers(20, 150))
        event_end = min(
            event_start + pd.Timedelta(minutes=duration),
            end,
        )

        equipment_id = None
        if event_type in {
            "equipment_fault",
            "maintenance",
            "consumable_rework",
        }:
            equipment_id = equipment.iloc[
                int(rng.integers(0, len(equipment)))
            ]["equipment_id"]

        work_order_id = None
        if event_type in {
            "quality_hold",
            "quality_stop",
            "material_shortage",
            "operator_absence",
        }:
            work_order_id = production.iloc[
                int(rng.integers(0, len(production)))
            ]["work_order_id"]

        rows.append(
            {
                "event_id": f"EV_{event_id:04d}",
                "event_type": event_type,
                "event_name": event_name,
                "start_time": event_start,
                "end_time": event_end,
                "severity": int(rng.integers(1, 4)),
                "affected_equipment_id": equipment_id,
                "affected_work_order_id": work_order_id,
                "affected_process_code": None,
                "event_status": "closed",
                "event_label": 1,
                "description": event_name,
            }
        )

    return pd.DataFrame(rows)


def generate_logistics(cfg: Config) -> pd.DataFrame:
    rng = np.random.default_rng(cfg.random_seed + 200)

    start = pd.Timestamp(cfg.start_time, tz="UTC")
    index = pd.date_range(
        start=start,
        periods=cfg.hours * 60 // cfg.telemetry_minutes,
        freq=f"{cfg.telemetry_minutes}min",
        tz="UTC",
    )

    materials = [
        ("RAW_SILICON", "原材料", "单晶硅毛坯"),
        ("RAW_CHALCO", "原材料", "硫系玻璃毛坯"),
        ("RAW_MICRO", "原材料", "微晶玻璃毛坯"),
        ("CONS_WHEEL", "辅料", "砂轮"),
        ("CONS_POLISH", "辅料", "抛光液"),
        ("CONS_COATING", "辅料", "镀膜靶材"),
    ]

    rows = []

    for timestamp in index:
        for material_id, material_type, material_name in materials:
            safety_stock = 30
            warehouse_qty = max(
                0,
                int(80 + rng.normal(0, 12)),
            )
            required_qty = int(rng.integers(5, 25))

            shortage = int(warehouse_qty < required_qty)
            warning = int(
                warehouse_qty < safety_stock
            )

            rows.append(
                {
                    "timestamp": timestamp,
                    "material_id": material_id,
                    "material_type": material_type,
                    "material_name": material_name,
                    "warehouse_qty": warehouse_qty,
                    "safety_stock": safety_stock,
                    "required_qty": required_qty,
                    "arrival_rate": round(
                        max(0, rng.normal(1.0, 0.15)), 3
                    ),
                    "shortage_flag": shortage,
                    "consumable_warning_flag": warning,
                    "supplier_delay_flag": int(
                        rng.random() < 0.01
                    ),
                }
            )

    return pd.DataFrame(rows)


def generate_telemetry(
    cfg: Config,
    equipment: pd.DataFrame,
    production: pd.DataFrame,
    events: pd.DataFrame,
) -> pd.DataFrame:
    rng = np.random.default_rng(cfg.random_seed + 300)

    start = pd.Timestamp(cfg.start_time, tz="UTC")
    index = pd.date_range(
        start=start,
        periods=cfg.hours * 60 // cfg.telemetry_minutes,
        freq=f"{cfg.telemetry_minutes}min",
        tz="UTC",
    )

    rows = []

    for _, eq in equipment.iterrows():
        eq_id = eq["equipment_id"]
        eq_type = eq["equipment_type"]

        eq_process = production[
            production["equipment_id"] == eq_id
        ]

        fault_intervals = events[
            events["affected_equipment_id"] == eq_id
        ]

        for timestamp in index:
            active = eq_process[
                (eq_process["actual_start"] <= timestamp)
                & (eq_process["actual_end"] >= timestamp)
            ]

            running = int(not active.empty)
            work_order_id = None
            workpiece_id = None
            process_code = None
            product_type = None

            if running:
                current = active.iloc[0]
                work_order_id = current["work_order_id"]
                workpiece_id = current["workpiece_id"]
                process_code = current["process_code"]
                product_type = current["product_type"]

            fault = fault_intervals[
                (fault_intervals["start_time"] <= timestamp)
                & (fault_intervals["end_time"] >= timestamp)
            ]

            alarm_code = None
            if not fault.empty:
                alarm_code = fault.iloc[0]["event_type"]
                running = 0

            rows.append(
                {
                    "timestamp": timestamp,
                    "equipment_id": eq_id,
                    "equipment_type": eq_type,
                    "product_type": product_type,
                    "work_order_id": work_order_id,
                    "workpiece_id": workpiece_id,
                    "process_code": process_code,
                    "running": running,
                    "spindle_speed_rpm": round(
                        max(
                            0,
                            rng.normal(
                                1800 if running else 0,
                                100,
                            ),
                        ),
                        2,
                    ),
                    "feed_rate_mm_min": round(
                        max(
                            0,
                            rng.normal(
                                60 if running else 0,
                                8,
                            ),
                        ),
                        2,
                    ),
                    "power_kw": round(
                        max(
                            0,
                            rng.normal(
                                18 if running else 1.5,
                                2.5,
                            ),
                        ),
                        3,
                    ),
                    "vibration_rms": round(
                        max(
                            0,
                            rng.normal(
                                0.5 if running else 0.2,
                                0.08,
                            ),
                        ),
                        4,
                    ),
                    "temperature_c": round(
                        rng.normal(
                            27 if running else 24,
                            1.0,
                        ),
                        3,
                    ),
                    "pressure_mpa": round(
                        max(
                            0,
                            rng.normal(
                                0.35 if running else 0.1,
                                0.04,
                            ),
                        ),
                        4,
                    ),
                    "vacuum_pa": round(
                        max(
                            0,
                            rng.normal(
                                0.002 if eq_type == "vacuum_coater"
                                else 0.1,
                                0.02,
                            ),
                        ),
                        6,
                    ),
                    "coolant_flow_l_min": round(
                        max(
                            0,
                            rng.normal(
                                12 if running else 0,
                                1.5,
                            ),
                        ),
                        3,
                    ),
                    "alarm_code": alarm_code,
                    "equipment_health_label": (
                        2 if alarm_code == "equipment_fault"
                        else 1 if alarm_code else 0
                    ),
                }
            )

    return pd.DataFrame(rows)


def build_minute_labels(
    cfg: Config,
    telemetry: pd.DataFrame,
    logistics: pd.DataFrame,
    events: pd.DataFrame,
) -> pd.DataFrame:
    index = pd.date_range(
        start=pd.Timestamp(cfg.start_time, tz="UTC"),
        periods=cfg.hours * 60 // cfg.telemetry_minutes,
        freq=f"{cfg.telemetry_minutes}min",
        tz="UTC",
    )

    result = pd.DataFrame(index=index)
    result.index.name = "timestamp"

    equipment_status = (
        telemetry.groupby("timestamp")["equipment_health_label"]
        .max()
        .reindex(index)
        .fillna(0)
        .astype(int)
    )

    running_count = (
        telemetry.groupby("timestamp")["running"]
        .sum()
        .reindex(index)
        .fillna(0)
    )

    logistics_status = (
        logistics.groupby("timestamp")
        .agg(
            shortage=("shortage_flag", "max"),
            warning=("consumable_warning_flag", "max"),
        )
        .reindex(index)
        .fillna(0)
    )

    result["equipment_state_label"] = equipment_status.to_numpy()
    result["logistics_state_label"] = np.select(
        [
            logistics_status["shortage"].to_numpy() > 0,
            logistics_status["warning"].to_numpy() > 0,
        ],
        [2, 1],
        default=0,
    )
    result["line_state_label"] = np.select(
        [
            running_count.to_numpy() == 0,
            equipment_status.to_numpy() >= 2,
            running_count.to_numpy() <= 1,
        ],
        [2, 3, 1],
        default=0,
    )

    result["process_state_label"] = (
        telemetry.groupby("timestamp")["running"]
        .max()
        .reindex(index)
        .fillna(0)
        .astype(int)
        .to_numpy()
    )

    result["anomaly_flag"] = 0
    result["event_type"] = None

    for _, event in events.iterrows():
        mask = (
            (result.index >= event["start_time"])
            & (result.index <= event["end_time"])
        )
        result.loc[mask, "anomaly_flag"] = 1
        result.loc[mask, "event_type"] = event["event_type"]

    result["quality_state_label"] = np.where(
        result["event_type"].eq("quality_hold"),
        2,
        0,
    )

    result["plan_state_label"] = np.select(
        [
            result["event_type"].eq("insert_order"),
            result["event_type"].eq("plan_change"),
            result["event_type"].eq("quality_hold"),
        ],
        [3, 1, 2],
        default=0,
    )

    result["personnel_state_label"] = np.where(
        result["event_type"].eq("operator_absence"),
        1,
        0,
    )

    result["facility_state_label"] = np.where(
        result["event_type"].eq("facility_fault"),
        1,
        0,
    )

    return result.reset_index()


def write_data_dictionary(output_dir: Path):
    if Document is None:
        print("未安装 python-docx，跳过 Word 文件生成。")
        return

    document = Document()
    document.add_heading("光学元件多品种产线数据字典", level=1)

    document.add_paragraph(
        "本数据集用于模拟单晶硅、硫系玻璃和微晶玻璃三类工件的生产过程，"
        "同时覆盖设备、工艺、质量、物流、计划、人员、设施和异常事件数据。"
    )

    tables = {
        "equipment_telemetry.csv": [
            ("timestamp", "UTC 时间"),
            ("equipment_id", "设备唯一编号"),
            ("equipment_type", "设备类型"),
            ("work_order_id", "当前工单"),
            ("workpiece_id", "当前工件"),
            ("process_code", "当前工序编码"),
            ("running", "设备是否运行"),
            ("spindle_speed_rpm", "主轴转速"),
            ("feed_rate_mm_min", "进给速度"),
            ("power_kw", "设备功率"),
            ("vibration_rms", "振动有效值"),
            ("temperature_c", "设备温度"),
            ("pressure_mpa", "压力"),
            ("vacuum_pa", "真空度"),
            ("coolant_flow_l_min", "冷却液流量"),
            ("alarm_code", "报警或异常编码"),
            ("equipment_health_label", "设备健康标签"),
        ],
        "production_records.csv": [
            ("work_order_id", "工单编号"),
            ("workpiece_id", "工件编号"),
            ("product_type", "产品类型"),
            ("process_sequence", "工序顺序"),
            ("process_code", "工序编码"),
            ("process_name", "工序名称"),
            ("equipment_id", "执行设备"),
            ("planned_start", "计划开始时间"),
            ("planned_end", "计划结束时间"),
            ("actual_start", "实际开始时间"),
            ("actual_end", "实际结束时间"),
            ("operator_id", "操作人员"),
            ("process_status", "工序状态"),
        ],
        "events.csv": [
            ("event_id", "事件编号"),
            ("event_type", "事件类型"),
            ("event_name", "事件名称"),
            ("start_time", "事件开始"),
            ("end_time", "事件结束"),
            ("severity", "严重程度"),
            ("affected_equipment_id", "受影响设备"),
            ("affected_work_order_id", "受影响工单"),
            ("event_label", "异常事件标签"),
        ],
        "minute_labels.csv": [
            ("process_state_label", "工艺状态"),
            ("equipment_state_label", "设备状态"),
            ("logistics_state_label", "物流物料状态"),
            ("line_state_label", "产线状态"),
            ("quality_state_label", "质量状态"),
            ("plan_state_label", "计划状态"),
            ("personnel_state_label", "人员状态"),
            ("facility_state_label", "设施环境状态"),
            ("anomaly_flag", "是否处于异常事件"),
            ("event_type", "当前异常类型"),
        ],
    }

    for filename, fields in tables.items():
        document.add_heading(filename, level=2)
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "字段"
        table.rows[0].cells[1].text = "说明"

        for field_name, description in fields:
            cells = table.add_row().cells
            cells[0].text = field_name
            cells[1].text = description

    document.add_heading("建模建议", level=2)
    document.add_paragraph(
        "建议以工单、工件和工序作为核心关联键；设备遥测采用 timestamp + "
        "equipment_id 关联；质量数据采用 workpiece_id + process_code 关联；"
        "异常事件采用 start_time、end_time 和影响对象关联。"
    )
    document.add_paragraph(
        "训练集、验证集和测试集应优先按照工单或时间进行划分，避免同一工件的"
        "相邻窗口同时出现在训练集和测试集中，从而造成数据泄漏。"
    )

    document.save(output_dir / "production_data_dictionary.docx")


def generate_all(cfg: Config):
    random.seed(cfg.random_seed)
    np.random.seed(cfg.random_seed)

    output_dir = Path(cfg.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    master = equipment_master(cfg)
    production, quality, plan, personnel = generate_work_orders(
        cfg, master
    )
    events = generate_events(cfg, master, production)
    logistics = generate_logistics(cfg)
    telemetry = generate_telemetry(
        cfg, master, production, events
    )
    labels = build_minute_labels(
        cfg, telemetry, logistics, events
    )

    dataframes = {
        "equipment_master": master,
        "production_records": production,
        "quality_inspection": quality,
        "production_plan": plan,
        "personnel_shift": personnel,
        "events": events,
        "logistics_material": logistics,
        "equipment_telemetry": telemetry,
        "minute_labels": labels,
    }

    for name, dataframe in dataframes.items():
        dataframe.to_csv(
            output_dir / f"{name}.csv",
            index=False,
            encoding="utf-8-sig",
        )

    write_data_dictionary(output_dir)

    print(f"数据已生成到：{output_dir.resolve()}")
    for name, dataframe in dataframes.items():
        print(f"{name:24s} {dataframe.shape}")


if __name__ == "__main__":
    generate_all(Config())
